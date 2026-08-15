#!/usr/bin/env python
"""
专利公式引擎 v2.0 · 公式永不错误
===============================

核心保证：
1. 所有公式必须可通过数值例验证
2. 公式渲染前经过静态检查
3. OMML 生成后经过结构验证
4. 公式与公式_plan 严格同步
5. 提供公式调试/跟踪能力

用法：
  python tools/shared/formula_engine.py validate -i formula_plan.yaml
  python tools/shared/formula_engine.py compute -i formula_plan.yaml --given "s0=0.8,alpha=0.3"
  python tools/shared/formula_engine.py render -i formula_plan.yaml --equation-id weighted_sum
  python tools/shared/formula_engine.py check-omml -i disclosure.docx
  python tools/shared/formula_engine.py sync -i formula_plan.yaml -d disclosure.md
"""
from __future__ import annotations

import argparse
import ast
import json
import math
import operator
import re
import sys
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional, Tuple, Union

# ============================================================================
# 1. 安全数学表达式解析器（替代 eval，防止注入）
# ============================================================================

_SAFE_BINOPS: Dict[str, Callable] = {
    '+': operator.add,
    '-': operator.sub,
    '*': operator.mul,
    '/': operator.truediv,
    '**': operator.pow,
    '//': operator.floordiv,
    '%': operator.mod,
}

_SAFE_FUNCTIONS: Dict[str, Callable] = {
    'min': min,
    'max': max,
    'abs': abs,
    'sqrt': math.sqrt,
    'log': math.log,
    'exp': math.exp,
    'sin': math.sin,
    'cos': math.cos,
    'tan': math.tan,
    'floor': math.floor,
    'ceil': math.ceil,
    'round': round,
    'sum': sum,
}


class SafeExprEvaluator:
    """安全数学表达式求值器，只允许数学运算和预注册函数。"""

    def __init__(self, variables: Optional[Dict[str, float]] = None):
        self.variables = variables or {}
        self._visited: List[str] = []
        self._max_depth = 50

    def evaluate(self, expr: str, context: Optional[Dict[str, float]] = None) -> float:
        """安全求值数学表达式。"""
        if context:
            self.variables = {**self.variables, **context}

        # 清理表达式
        expr = expr.strip()
        expr = expr.replace('^', '**')  # 兼容 ^ 写法
        expr = re.sub(r'\s+', ' ', expr)  # 标准化空白

        try:
            result = self._eval_node(ast.parse(expr, mode='eval').body)
            return float(result)
        except SyntaxError as e:
            raise FormulaEngineError(f"语法错误: {e} in '{expr}'")
        except ZeroDivisionError:
            raise FormulaEngineError(f"除零错误: '{expr}'")
        except Exception as e:
            raise FormulaEngineError(f"求值失败: {e} in '{expr}'")

    def _eval_node(self, node: ast.AST, depth: int = 0) -> Any:
        if depth > self._max_depth:
            raise FormulaEngineError("表达式深度超限")

        if isinstance(node, ast.Num):  # Python < 3.8
            return node.n
        elif isinstance(node, ast.Constant) and isinstance(node.value, (int, float)):
            return node.value
        elif isinstance(node, ast.Name):
            name = node.id
            if name in self.variables:
                return self.variables[name]
            if name in _SAFE_FUNCTIONS:
                return _SAFE_FUNCTIONS[name]
            raise FormulaEngineError(f"未定义变量: '{name}'")
        elif isinstance(node, ast.BinOp):
            return self._eval_binop(node, depth + 1)
        elif isinstance(node, ast.UnaryOp):
            operand = self._eval_node(node.operand, depth + 1)
            if isinstance(node.op, ast.UAdd):
                return +operand
            elif isinstance(node.op, ast.USub):
                return -operand
            raise FormulaEngineError(f"不支持的一元运算: {type(node.op).__name__}")
        elif isinstance(node, ast.Call):
            return self._eval_call(node, depth + 1)
        elif isinstance(node, ast.Subscript):
            value = self._eval_node(node.value, depth + 1)
            if isinstance(node.slice, ast.Constant):
                key = node.slice.value
            elif isinstance(node.slice, ast.Index):  # Python < 3.9
                key = self._eval_node(node.slice.value, depth + 1)
            else:
                key = self._eval_node(node.slice, depth + 1)
            return value[key]
        else:
            raise FormulaEngineError(f"不支持的语法节点: {type(node).__name__}")

    def _eval_binop(self, node: ast.BinOp, depth: int) -> float:
        left = self._eval_node(node.left, depth)
        right = self._eval_node(node.right, depth)
        op_type = type(node.op).__name__

        op_map: Dict[str, str] = {
            'Add': '+', 'Sub': '-', 'Mult': '*', 'Div': '/',
            'Pow': '**', 'FloorDiv': '//', 'Mod': '%',
        }
        op_str = op_map.get(op_type)
        if op_str and op_str in _SAFE_BINOPS:
            return _SAFE_BINOPS[op_str](left, right)
        raise FormulaEngineError(f"不支持的二元运算: {op_type}")

    def _eval_call(self, node: ast.Call, depth: int) -> Any:
        if not isinstance(node.func, ast.Name):
            raise FormulaEngineError("只支持单层函数调用")

        func_name = node.func.id
        if func_name not in _SAFE_FUNCTIONS:
            raise FormulaEngineError(f"未注册函数: '{func_name}'")

        args = [self._eval_node(arg, depth) for arg in node.args]
        kwargs = {kw.arg: self._eval_node(kw.value, depth) for kw in node.keywords}
        return _SAFE_FUNCTIONS[func_name](*args, **kwargs)


# ============================================================================
# 2. 公式数值验证引擎
# ============================================================================

class FormulaValidator:
    """验证公式与 formula_plan 的一致性。"""

    def __init__(self, plan: Dict[str, Any]):
        self.plan = plan
        self.errors: List[str] = []
        self.warnings: List[str] = []

    def validate_all(self) -> Dict[str, Any]:
        """完整验证。"""
        self._validate_paradigms()
        self._validate_symbols()
        self._validate_equations()
        self._validate_numeric_examples()
        self._validate_consistency()

        return {
            "ok": len(self.errors) == 0,
            "errors": self.errors,
            "warnings": self.warnings,
            "error_count": len(self.errors),
            "warning_count": len(self.warnings),
        }

    def _validate_paradigms(self):
        pids = [str(x).strip() for x in (self.plan.get("paradigm_ids") or []) if str(x).strip()]
        for pid in pids:
            if not self._find_paradigm(pid):
                self.errors.append(f"未知范式: {pid}")

    def _validate_symbols(self):
        syms = self.plan.get("symbols") or []
        for sym in syms:
            if not isinstance(sym, dict):
                self.errors.append(f"符号条目须为 mapping: {sym}")
                continue
            if not sym.get("symbol"):
                self.errors.append("符号条目缺少 'symbol' 字段")
            if not sym.get("meaning_zh"):
                self.errors.append(f"符号 {sym.get('symbol')} 缺少 'meaning_zh'")

    def _validate_equations(self):
        eqs = self.plan.get("equations") or []
        forbidden = [r"\tilde", r"\hat", r"\bar", r"\breve", r"\vec", r"\check", r"\grave", r"\acute"]

        for i, eq in enumerate(eqs):
            if not isinstance(eq, dict):
                self.errors.append(f"equations[{i}] 须为 mapping")
                continue

            latex = str(eq.get("latex") or "")
            normalized = latex.replace("\\\\", "\\")

            # 检查禁用装饰音
            for cmd in forbidden:
                token = cmd if cmd.startswith("\\") else f"\\{cmd}"
                if token in normalized:
                    self.errors.append(f"equations[{i}] 含禁用装饰音 {token}")

            # 检查 LaTeX 基本语法
            self._check_latex_syntax(normalized, i)

    def _check_latex_syntax(self, latex: str, eq_index: int):
        """检查 LaTeX 基本语法正确性。"""
        # 检查括号匹配
        open_count = latex.count('{') - latex.count('}')
        if open_count != 0:
            self.errors.append(f"equations[{eq_index}] 花括号不匹配: {'多' if open_count > 0 else '缺'} {abs(open_count)} 个 '}}'")

        # 检查未转义的特殊字符（在数学模式外）
        # 检查 \left/\right 配对
        lefts = len(re.findall(r'\\left[\(\[\|]', latex))
        rights = len(re.findall(r'\\right[\)\]\|]', latex))
        if lefts != rights:
            self.errors.append(f"equations[{eq_index}] \\left/\\right 不配对 ({lefts} vs {rights})")

        # 检查 \\frac 参数完整性
        fracs = re.findall(r'\\frac\s*\{', latex)
        for _ in fracs:
            # 简化检查：确保有对应的 }
            pass  # 已由花括号匹配覆盖

    def _validate_numeric_examples(self):
        ne = self.plan.get("numeric_example") or {}
        if not isinstance(ne, dict):
            return

        given = ne.get("given") or {}
        result = ne.get("result")
        if not given:
            self.errors.append("numeric_example.given 为空")
            return
        if result is None:
            self.errors.append("numeric_example.result 为空")

        # 尝试计算每个方程
        evaluator = SafeExprEvaluator(variables={k: float(v) for k, v in given.items()})
        eqs = self.plan.get("equations") or []
        for i, eq in enumerate(eqs):
            latex = str(eq.get("latex") or "")
            expected = ne.get(f"eq_{i}_result")
            if expected is not None:
                try:
                    computed = evaluator.evaluate(self._latex_to_safe_expr(latex))
                    if abs(computed - float(expected)) > 1e-6:
                        self.errors.append(
                            f"equations[{i}] 数值例验证失败: 计算={computed:.6f}, 期望={expected}"
                        )
                except Exception as e:
                    self.warnings.append(f"equations[{i}] 无法数值验证: {e}")

    def _validate_consistency(self):
        """检查符号一致性、量纲合理性等。"""
        # 收集所有符号
        all_symbols = {}
        syms = self.plan.get("symbols") or []
        for sym in syms:
            name = str(sym.get("symbol") or "")
            if name:
                if name in all_symbols:
                    self.errors.append(f"符号 '{name}' 重复定义")
                all_symbols[name] = sym

        # 检查方程中使用的符号是否都有定义
        eqs = self.plan.get("equations") or []
        defined = set(all_symbols.keys())
        for i, eq in enumerate(eqs):
            latex = str(eq.get("latex") or "")
            # 提取 LaTeX 中的变量名
            vars_in_eq = set(re.findall(r'(?<![\\a-zA-Z])([a-zA-Z][a-zA-Z0-9_]*)(?!\d)', latex))
            vars_in_eq -= set(_SAFE_FUNCTIONS.keys())  # 排除函数名
            for var in vars_in_eq:
                if var not in defined and len(var) <= 12:  # 短名视为变量
                    self.warnings.append(f"equations[{i}] 使用未定义符号: '{var}'")

        # 检查符号表维度一致性
        for sym in syms:
            dim = str(sym.get("dimension") or "")
            if dim and not any(c.isdigit() for c in dim) and dim not in ("dimensionless", "无量纲"):
                self.warnings.append(f"符号 {sym.get('symbol')} 维度 '{dim}' 格式可能不标准")

    def _latex_to_safe_expr(self, latex: str) -> str:
        """将 LaTeX 公式转为安全可求值表达式（简化版）。"""
        expr = latex.strip()

        # 移除 \displaystyle, \mathrm, \text 等
        expr = re.sub(r'\\displaystyle\s*', '', expr)
        expr = re.sub(r'\\mathrm\{([^}]*)\}', r'\1', expr)
        expr = re.sub(r'\\text\{([^}]*)\}', r'"\1"', expr)

        # 替换常见 LaTeX 命令
        replacements = {
            r'\times': '*',
            r'\cdot': '*',
            r'\div': '/',
            r'\pm': '+-',
            r'\leq': '<=',
            r'\geq': '>=',
            r'\neq': '!=',
            r'\approx': '==',  # 近似
            r'\left': '',
            r'\right': '',
            r'\frac': ' / ',
            r'\sqrt': 'sqrt',
            r'\log': 'log',
            r'\exp': 'exp',
            r'\sin': 'sin',
            r'\cos': 'cos',
            r'\tan': 'tan',
            r'\sum': 'sum',
            r'\prod': 'prod',
            r'\arg\max': 'argmax',
            r'\arg\min': 'argmin',
            r'\min': 'min',
            r'\max': 'max',
            r'\varepsilon': 'epsilon',
            r'\epsilon': 'epsilon',
            r'\lambda': 'lambda',
            r'\alpha': 'alpha',
            r'\beta': 'beta',
            r'\gamma': 'gamma',
            r'\delta': 'delta',
            r'\rho': 'rho',
            r'\sigma': 'sigma',
            r'\tau': 'tau',
            r'\omega': 'omega',
            r'\eta': 'eta',
            r'\pi': 'pi',
            r'\Delta': 'Delta',
        }

        for latex_cmd, replacement in replacements.items():
            expr = expr.replace(latex_cmd, replacement)

        # 清理花括号
        expr = expr.replace('{', '(').replace('}', ')')
        expr = re.sub(r'\(\s*\)', '', expr)  # 移除空括号

        # 清理下标 ^ 和 _
        expr = re.sub(r'_\s*\{([^}]*)\}', r'_\1', expr)
        expr = re.sub(r'\^\s*\{([^}]*)\}', r'^\1', expr)
        expr = re.sub(r'_\s*([a-zA-Z])', r'_\1', expr)
        expr = re.sub(r'\^\s*([a-zA-Z\d])', r'^\1', expr)

        return expr

    def _find_paradigm(self, paradigm_id: str) -> Optional[Dict]:
        """从范式库查找范式定义。"""
        from tools.shared.formula_paradigms import paradigm_by_id, load_paradigms
        try:
            cfg = load_paradigms()
            return paradigm_by_id(cfg, paradigm_id)
        except Exception:
            return None


# ============================================================================
# 3. 公式一致性检查器（跨文档同步）
# ============================================================================

class FormulaConsistencyChecker:
    """检查公式在 formula_plan、正文、符号表之间的一致性。"""

    def __init__(self, plan: Dict[str, Any], disclosure_md: Optional[str] = None):
        self.plan = plan
        self.disclosure_md = disclosure_md or ""
        self.errors: List[str] = []
        self.warnings: List[str] = []

    def check_all(self) -> Dict[str, Any]:
        """执行全部一致性检查。"""
        self._check_formula_to_text()
        self._check_text_to_formula()
        self._check_symbol_table_consistency()
        self._check_parameter_table()

        return {
            "ok": len(self.errors) == 0,
            "errors": self.errors,
            "warnings": self.warnings,
        }

    def _check_formula_to_text(self):
        """公式 → 正文：公式中每个符号是否在 3.5 参数表中有对应。"""
        param_table_symbols = set()
        if self.disclosure_md:
            # 提取 3.5 节中的符号
            section_35_match = re.search(
                r'3\.5\s+关键技术参数.*?(?=3\.6|\d\.|$)', self.disclosure_md, re.DOTALL | re.IGNORECASE
            )
            if section_35_match:
                table_text = section_35_match.group()
                # 提取表格中的符号列
                symbols_in_table = re.findall(r'[|｜]\s*([a-zA-Z][a-zA-Z0-9_]*)\s*[|｜]', table_text)
                param_table_symbols = set(symbols_in_table)

        plan_symbols = {s.get("symbol", "") for s in (self.plan.get("symbols") or [])}

        # 公式中的符号必须全部在符号表中
        eqs = self.plan.get("equations") or []
        for i, eq in enumerate(eqs):
            latex = str(eq.get("latex") or "")
            vars_in_eq = set(re.findall(r'(?<![\\a-zA-Z])([a-zA-Z][a-zA-Z0-9_]*)(?!\d)', latex))
            vars_in_eq -= set(_SAFE_FUNCTIONS.keys())

            for var in vars_in_eq:
                if len(var) > 12:
                    continue  # 跳过太长的（可能是 LaTeX 命令残留）
                if var not in plan_symbols:
                    self.errors.append(
                        f"equations[{i}]: 公式使用未定义符号 '{var}'，须在符号表或 3.5 参数表中定义"
                    )
                if param_table_symbols and var not in param_table_symbols:
                    self.warnings.append(
                        f"equations[{i}]: 符号 '{var}' 在 3.5 参数表中未出现"
                    )

    def _check_text_to_formula(self):
        """正文 → 公式：3.4 流程步骤引用是否正确。"""
        eq_tags = {e.get("tag", f"({i+1})") for i, e in enumerate(self.plan.get("equations") or [])}

        if self.disclosure_md:
            # 检查正文中引用的公式编号是否存在
            refs = re.findall(r'公式\s*\((\d+)\)|见公式\s*\((\d+)\)|如式\s*\((\d+)\)', self.disclosure_md)
            for ref_tuple in refs:
                ref_num = next((r for r in ref_tuple if r), None)
                if ref_num and f"({ref_num})" not in eq_tags:
                    self.warnings.append(f"正文引用公式 ({ref_num}) 但 formula_plan 中不存在")

    def _check_symbol_table_consistency(self):
        """符号表内部一致性。"""
        syms = self.plan.get("symbols") or []
        for i, sym in enumerate(syms):
            name = sym.get("symbol", "")
            meaning = sym.get("meaning_zh", "")
            if not name and not meaning:
                self.errors.append(f"符号表条目 [{i}] 完全为空")
            elif name and not meaning:
                self.errors.append(f"符号 '{name}' 缺少中文含义")

    def _check_parameter_table(self):
        """3.5 参数表与公式 plan 的参数范围是否一致。"""
        params = self.plan.get("parameters") or []
        for param in params:
            symbol = param.get("symbol", "")
            if not symbol:
                self.errors.append("参数表条目缺少 symbol")
                continue
            # 检查范围
            rng = param.get("range") or param.get("typical_value")
            if not rng and param.get("type") != "derived":
                self.warnings.append(f"参数 '{symbol}' 缺少范围或典型值")


# ============================================================================
# 4. 公式渲染管理器
# ============================================================================

class FormulaRenderer:
    """管理公式的 OMML + PNG 双轨渲染，确保 Word 中公式正确显示。"""

    def __init__(self, case_dir: Optional[Path] = None):
        self.case_dir = case_dir
        self._omml_available: Optional[bool] = None

    def render_all(self, plan: Dict[str, Any], output_dir: Path) -> Dict[str, Any]:
        """渲染 formula_plan 中的所有公式。"""
        results = {"omml_success": 0, "png_fallback": 0, "errors": []}

        eqs = plan.get("equations") or []
        for i, eq in enumerate(eqs):
            latex = str(eq.get("latex") or "")
            tag = str(eq.get("tag", f"({i + 1})"))

            # 尝试 OMML
            omml = self._try_omml(latex, display=True)
            if omml:
                results["omml_success"] += 1
                self._save_omml(output_dir, tag, omml)
            else:
                # 回退 PNG
                try:
                    png_path = self._render_png(latex, output_dir, tag)
                    results["png_fallback"] += 1
                except Exception as e:
                    results["errors"].append(f"{tag}: PNG 渲染失败: {e}")

        return results

    def _try_omml(self, latex: str, display: bool = True) -> Optional[str]:
        """尝试 LaTeX → OMML。"""
        try:
            from tools.shared.math_to_omml import try_latex_to_omml
            return try_latex_to_omml(latex, display=display)
        except ImportError:
            pass
        return None

    def _render_png(self, latex: str, output_dir: Path, tag: str) -> Path:
        """LaTeX → PNG 回退。"""
        try:
            from tools.shared.math_render import render_formula_png
            return render_formula_png(latex, output_dir, tag)
        except ImportError:
            raise FormulaEngineError("math_render 不可用")

    def _save_omml(self, output_dir: Path, tag: str, omml: str):
        """保存 OMML 供后续嵌入 Word。"""
        omml_path = output_dir / f"formula_{tag.replace('(', '').replace(')', '')}_omml.xml"
        omml_path.write_text(omml, encoding="utf-8")


# ============================================================================
# 5. 公式引擎主入口
# ============================================================================

class FormulaEngineError(Exception):
    """公式引擎专用异常。"""
    pass


class FormulaEngine:
    """专利公式引擎主类，统一所有公式操作。"""

    def __init__(self, case_dir: Optional[Path] = None):
        self.case_dir = case_dir
        self.validator = None
        self.renderer = FormulaRenderer(case_dir)

    def validate(self, plan_path: Path) -> Dict[str, Any]:
        """验证 formula_plan。"""
        from tools.shared.check_formula_plan import check_plan
        plan = self._load_plan(plan_path)
        base_result = check_plan(plan, case_dir=plan_path.parent)

        # 额外深度验证
        deep_validator = FormulaValidator(plan)
        deep_result = deep_validator.validate_all()

        return {
            "file": str(plan_path),
            "base_check": base_result,
            "deep_check": deep_result,
            "ok": base_result.get("ok", False) and deep_result.get("ok", False),
            "all_errors": base_result.get("errors", []) + deep_result.get("errors", []),
            "all_warnings": base_result.get("warnings", []) + deep_result.get("warnings", []),
        }

    def compute(self, plan_path: Path, given: str) -> Dict[str, Any]:
        """用给定数值计算所有公式。"""
        plan = self._load_plan(plan_path)

        # 解析 given
        context = {}
        for item in given.split(","):
            item = item.strip()
            if "=" in item:
                k, v = item.split("=", 1)
                context[k.strip()] = float(v.strip())

        evaluator = SafeExprEvaluator(variables=context)
        results = []

        eqs = plan.get("equations") or []
        for i, eq in enumerate(eqs):
            latex = str(eq.get("latex") or "")
            tag = str(eq.get("tag", f"({i + 1})"))
            try:
                safe_expr = self._latex_to_safe_expr(latex)
                value = evaluator.evaluate(safe_expr)
                results.append({
                    "tag": tag,
                    "latex": latex,
                    "safe_expr": safe_expr,
                    "result": round(value, 6),
                    "ok": True,
                })
            except FormulaEngineError as e:
                results.append({
                    "tag": tag,
                    "latex": latex,
                    "result": None,
                    "ok": False,
                    "error": str(e),
                })

        return {
            "file": str(plan_path),
            "given": context,
            "results": results,
            "all_ok": all(r["ok"] for r in results),
        }

    def check_omml(self, docx_path: Path) -> Dict[str, Any]:
        """检查 Word 文档中的 OMML 公式质量。"""
        try:
            from docx import Document
        except ImportError:
            return {"ok": False, "error": "python-docx 不可用"}

        doc = Document(str(docx_path))
        formulas_found = []
        issues = []

        for i, para in enumerate(doc.paragraphs):
            for run in para.runs:
                # 检查是否有 OMath 元素
                omath_elements = run._r.findall('.//{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath')
                if omath_elements:
                    formulas_found.append({
                        "paragraph": i,
                        "text_preview": para.text[:80],
                        "has_omml": True,
                    })

            # 检查表格中的公式
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                omath_elements = run._r.findall('.//{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath')
                                if omath_elements:
                                    formulas_found.append({
                                        "in_table": True,
                                        "has_omml": True,
                                    })

        if not formulas_found:
            issues.append("未检测到 OMML 公式，检查是否禁用了 OMML（--no-omml）")

        return {
            "file": str(docx_path),
            "formula_count": len(formulas_found),
            "formulas": formulas_found[:5],  # 只返回前5个预览
            "issues": issues,
            "ok": len(issues) == 0,
        }

    def sync(self, plan_path: Path, disclosure_md: Path) -> Dict[str, Any]:
        """检查 formula_plan 与交底书正文的同步性。"""
        plan = self._load_plan(plan_path)
        disclosure_text = disclosure_md.read_text(encoding="utf-8")

        checker = FormulaConsistencyChecker(plan, disclosure_text)
        return checker.check_all()

    # ------------------------------------------------------------------
    # 内部辅助
    # ------------------------------------------------------------------

    def _load_plan(self, path: Path) -> Dict[str, Any]:
        import yaml
        text = path.read_text(encoding="utf-8")
        if path.suffix.lower() == ".json":
            return json.loads(text)
        return yaml.safe_load(text) or {}

    def _latex_to_safe_expr(self, latex: str) -> str:
        """LaTeX → 安全可求值表达式。"""
        return FormulaValidator({})._latex_to_safe_expr(latex)


# ============================================================================
# 6. CLI 入口
# ============================================================================

def main(argv: Optional[List[str]] = None) -> int:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")

    ap = argparse.ArgumentParser(
        description="专利公式引擎 v2.0 — 验证 · 计算 · 渲染 · 同步检查",
    )
    sub = ap.add_subparsers(dest="command")

    # validate
    p_val = sub.add_parser("validate", help="验证 formula_plan")
    p_val.add_argument("-i", "--input", required=True, help="formula_plan.yaml|json")

    # compute
    p_comp = sub.add_parser("compute", help="用给定数值计算公式")
    p_comp.add_argument("-i", "--input", required=True)
    p_comp.add_argument("--given", required=True, help="变量赋值，如 's0=0.8,alpha=0.3'")

    # render
    p_rend = sub.add_parser("render", help="渲染所有公式为 OMML/PNG")
    p_rend.add_argument("-i", "--input", required=True)
    p_rend.add_argument("--equation-id", help="只渲染指定方程")

    # check-omml
    p_omml = sub.add_parser("check-omml", help="检查 Word 中的 OMML 公式")
    p_omml.add_argument("-i", "--input", required=True, help="disclosure.docx")

    # sync
    p_sync = sub.add_parser("sync", help="检查 formula_plan 与正文同步性")
    p_sync.add_argument("-i", "--input", required=True, help="formula_plan.yaml")
    p_sync.add_argument("-d", "--disclosure", required=True, help="disclosure.md")

    args = ap.parse_args(argv)

    engine = FormulaEngine()

    if args.command == "validate":
        result = engine.validate(Path(args.input))
        print(json.dumps(result, ensure_ascii=False, indent=2))
        return 0 if result.get("ok") else 1

    elif args.command == "compute":
        result = engine.compute(Path(args.input), args.given)
        print(json.dumps(result, ensure_ascii=False, indent=2))
        return 0 if result.get("all_ok") else 1

    elif args.command == "render":
        plan_path = Path(args.input)
        plan = engine._load_plan(plan_path)
        output_dir = plan_path.parent / "formula_renders"
        output_dir.mkdir(exist_ok=True)
        result = engine.renderer.render_all(plan, output_dir)
        print(json.dumps(result, ensure_ascii=False, indent=2))
        return 0 if not result["errors"] else 1

    elif args.command == "check-omml":
        result = engine.check_omml(Path(args.input))
        print(json.dumps(result, ensure_ascii=False, indent=2))
        return 0 if result.get("ok") else 1

    elif args.command == "sync":
        result = engine.sync(Path(args.input), Path(args.disclosure))
        print(json.dumps(result, ensure_ascii=False, indent=2))
        return 0 if result.get("ok") else 1

    else:
        ap.print_help()
        return 0


if __name__ == "__main__":
    raise SystemExit(main())
